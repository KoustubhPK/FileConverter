import os, io
from django.db.models import Q
import base64
import uuid
import datetime
import math
from PIL import Image
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth import authenticate, login
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.forms import PasswordResetForm
from django.core.mail import EmailMultiAlternatives
from django.utils.html import strip_tags

from django.http import HttpResponse
from django.utils import timezone

from PyPDF2 import PdfReader
import PyPDF2

from .forms import CustomUserCreationForm, ForgotPasswordForm
from PIL import Image
from django.conf import settings

from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from docx import Document
#from docx2pdf import convert
from docx.shared import Inches, Pt
from django.urls import reverse
from django.core.paginator import Paginator

from django.http import HttpResponse, JsonResponse
from django.template.loader import get_template
from xhtml2pdf import  pisa

from django.views.generic.edit import CreateView
from django.views.generic.detail import DetailView

from .forms import PdfTextForm
from .models import PdfTextModel, DeletedPdf
from django.views.generic import View
from django.shortcuts import get_object_or_404, redirect

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from django.contrib.auth.decorators import user_passes_test
from .utils import get_media_files

from django.contrib.sites.shortcuts import get_current_site
from django.template.loader import render_to_string
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes
from django.contrib.auth.tokens import default_token_generator

import logging
from django.contrib.auth import logout

from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.core.exceptions import ValidationError
from django.core.exceptions import ObjectDoesNotExist

from django.urls import reverse_lazy
from django.contrib.auth import views as auth_views

from zipfile import ZipFile
import fitz

# Create your views here.

def index(request):
    return render(request, 'main.html')

def features(request):
    return render(request, 'features.html')

def pricing(request):
    return render(request, 'pricing.html')

def about(request):
    return render(request, 'about.html')

def change_password(request):
    return render(request, 'changepassword.html')
    
def user_login(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            
            # Authenticate the user
            user = authenticate(request, username=username, password=password)
            
            if user is not None:
                if user.is_active:
                    # Log the user in
                    login(request, user)
                    messages.success(request, "You've been logged in.")
                    return redirect('home')
                else:
                    messages.error(request, "Your account is not activated.")
            else:
                messages.error(request, "Invalid username or password.")
    else:
        form = AuthenticationForm()

    return render(request, 'login.html', {'form': form})

logger = logging.getLogger(__name__)

def user_signup(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            logger.debug(f"User created: {user}")
            email = form.cleaned_data.get('email')

            # Log the user in
            # login(request, user)

            # Send email verification link
            current_site = get_current_site(request)
            mail_subject = 'Activate your account'
            message = render_to_string('pdfgo/account_activation_email.html', {
                'user': user,
                'domain': current_site.domain,
                'uid': urlsafe_base64_encode(force_bytes(user.pk)),
                'token': default_token_generator.make_token(user),
            })
            
            # Remove HTML tags from the email message
            plain_text_message = strip_tags(message)

            email = EmailMultiAlternatives(mail_subject, plain_text_message, to=[email])
            email.attach_alternative(message, "text/html")  # Attach the HTML version

            # Send the email
            email.send()

            # Redirect to a success page after successful signup (change 'home' to your desired URL)
            messages.success(request, "Account created successfully.")
            return redirect('home')
        else:
            messages.error(request, 'Something went wrong.')
    else:
        # Use an empty UserCreationForm for initial rendering of the form
        form = CustomUserCreationForm()

    return render(request, 'signup.html', {'form': form})

User = get_user_model()

def account_activation(request, uidb64, token):
    try:
        uid = urlsafe_base64_decode(uidb64).decode()
        user = User.objects.get(pk=uid)
    except (TypeError, ValueError, OverflowError, ValidationError, User.DoesNotExist):
        user = None

    if user and default_token_generator.check_token(user, token):
        # Activate the user account
        user.is_active = True
        user.save()
        return render(request, 'pdfgo/account_activation_success.html', {'user': user})
    else:
        return render(request, 'pdfgo/account_activation_failure.html')

@login_required(login_url='login')
def user_logout(request):
    logout(request)
    messages.success(request, "You've been logged out.")
    return redirect('login')

def forgot_password(request):
    if request.method == 'POST':
        form = ForgotPasswordForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data.get('email')
            
            try:
                user = User.objects.get(email=email)
            except ObjectDoesNotExist:
                messages.error(request, 'This email address is not associated with any account. Please try again.')
                return render(request, 'pdfgo/forgot_password.html')

            # Generate a one-time use link for password reset
            current_site = get_current_site(request)
            uid = urlsafe_base64_encode(force_bytes(user.pk))
            token = default_token_generator.make_token(user)

            # Send email for password reset
            mail_subject = 'Reset Your Password'
            message_html = render_to_string('pdfgo/password_reset_email.html', {
                'user': user,
                'domain': current_site.domain,
                'uid': uid,
                'token': token,
            })
            text_message = strip_tags(message_html)  # Extract plain text version from HTML content

            # Create an EmailMultiAlternatives object
            email_message = EmailMultiAlternatives(mail_subject, text_message, to=[email])
            email_message.attach_alternative(message_html, "text/html")  # Attach HTML content

            email_message.send()

            messages.success(request, "An email has been sent with instructions to reset your password.")
            return redirect('login')
        else:
            messages.error(request, 'Invalid email. Please try again.')

    else:
        form = ForgotPasswordForm()

    return render(request, 'pdfgo/forgot_password.html', {'form': form})

# Define a custom test function to check if the user is an admin
def is_admin(user):
    return user.is_authenticated and user.is_staff

def convert_size(size_bytes):
    # Function to convert bytes to a human-readable format (KB, MB, GB, etc.)
    if size_bytes == 0:
        return "0 B"
    size_name = ("B", "KB", "MB", "GB", "TB", "PB", "EB", "ZB", "YB")
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_name[i]}"

def get_media_files():
    # Define the directory where your media files are stored
    media_directory = 'media'

    # Initialize an empty list to store media file information
    media_files = []

    # Iterate through the files in the media directory
    for filename in os.listdir(media_directory):
        file_path = os.path.join(media_directory, filename)
        
        # Check if it's a file and not a directory
        if os.path.isfile(file_path):
            # Get file information, including created and modified timestamps and file size
            created_at = datetime.datetime.fromtimestamp(os.path.getctime(file_path))
            modified_at = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
            file_size = os.path.getsize(file_path)
            
            # Convert file size to a human-readable format
            file_size_readable = convert_size(file_size)
            
            # Add the file information to the list
            media_files.append({
                'name': filename,
                'created_at': created_at,
                'modified_at': modified_at,
                'file_size': file_size_readable,
            })

    return media_files

# Use the user_passes_test decorator to restrict access to admin users
@user_passes_test(is_admin)
def media_files(request):
    # Get the current page number from the GET request
    page_number = request.GET.get('page')

    # Get the selected per-page value from the GET request, defaulting to 5 items per page
    per_page = int(request.GET.get('per_page', 5))

    # Retrieve all media files
    media_files = get_media_files()

    # Sorting
    sort_by = request.GET.get('sort')
    descending = request.GET.get('desc')
    
    reverse_order = descending == '1' if descending else False
    
    if sort_by == 'size':
        # Convert file sizes to integers for sorting
        media_files = sorted(media_files, key=lambda x: int(x.get(sort_by, 0)), reverse=reverse_order)
    else:
        media_files = sorted(media_files, key=lambda x: x.get(sort_by, ''), reverse=reverse_order)

    # Paginate the media files with the new per-page value
    paginator = Paginator(media_files, per_page)
    
    # Ensure that the current page number is within the valid range
    if page_number:
        page_number = max(1, min(int(page_number), paginator.num_pages))
    else:
        page_number = 1
    
    page_obj = paginator.get_page(page_number)

    # Get the total count of media files
    total_count = paginator.count

    return render(request, 'pdfgo/media-files.html', {
        'media_files': page_obj,
        'total_count': total_count,
        'paginator': paginator,
        'sort_by': sort_by,
        'ascending': not reverse_order,
    })

def custom_page_not_found_view(request, exception):
    return render(request, "pdfgo/404.html", {})

def custom_error_view(request, exception=None):
    return render(request, "pdfgo/500.html", {})

def custom_permission_denied_view(request, exception=None):
    return render(request, "pdfgo/403.html", {})

def custom_bad_request_view(request, exception=None):
    return render(request, "pdfgo/400.html", {})

class CustomPasswordResetView(auth_views.PasswordResetView):
    template_name = 'custom_password_reset_form.html'
    email_template_name = 'password_reset_email.html'
    subject_template_name = 'password_reset_subject.txt'
    success_url = reverse_lazy('custom_password_reset_done')

class CustomPasswordResetDoneView(auth_views.PasswordResetDoneView):
    template_name = 'custom_password_reset_done.html'

class CustomPasswordResetConfirmView(auth_views.PasswordResetConfirmView):
    template_name = 'custom_password_reset_confirm.html'
    success_url = reverse_lazy('custom_password_reset_complete')

class CustomPasswordResetCompleteView(auth_views.PasswordResetCompleteView):
    template_name = 'custom_password_reset_complete.html'

@login_required(login_url='login')
def profile(request):
    user = request.user
    user_profiles = PdfTextModel.objects.filter(user=user)
    active_pdfs = PdfTextModel.objects.filter(user=request.user, deleted=False)
    modified_pdf = PdfTextModel.objects.filter(user=request.user)
    deleted_pdfs = DeletedPdf.objects.filter(user=request.user)

    active_pdf_count = active_pdfs.count()
    deleted_pdf_count = deleted_pdfs.count()
    modified_pdf_count = modified_pdf.count()
    total_pdf = active_pdf_count + deleted_pdf_count

    context = {
        'user_profiles': user_profiles,
        'deleted_pdf_count': deleted_pdf_count,
        'modified_pdf_count': modified_pdf_count,
        'total_pdf': total_pdf
    }

    return render(request, 'profile.html', context)

@login_required(login_url='login')
def archive(request):
    try:
        search_query = request.GET.get('q')
        sort_by = request.GET.get('sort', 'created_at')
        descending = request.GET.get('desc')

        pdf_queryset = PdfTextModel.objects.filter(user=request.user)  # Filter by the current user

        # Perform search based on the 'name' field
        if search_query:
            pdf_queryset = pdf_queryset.filter(Q(name__icontains=search_query))

        # Sort the queryset based on the selected field
        if sort_by == 'name':
            pdf_queryset = pdf_queryset.order_by('-name' if descending else 'name')
        elif sort_by == 'created_at':
            pdf_queryset = pdf_queryset.order_by('-created_at' if descending else 'created_at')
        elif sort_by == 'modified_at':
            pdf_queryset = pdf_queryset.order_by('-modified_at' if descending else 'modified_at')

        # Get the per_page value from the request.GET dictionary
        per_page = int(request.GET.get('per_page', 5))  # Default to 5 if not provided

        # Create the paginator object with the dynamically selected per_page value
        paginator = Paginator(pdf_queryset, per_page)

        # Get the page number from the request.GET dictionary
        page_number = request.GET.get('page')

        # Get the page object based on the page number
        page_obj = paginator.get_page(page_number)

        context = {
            'user_profiles': page_obj,  # Pass the paginated object to the template
            'sort_by': sort_by,
            'ascending': not descending,  # Invert the descending flag to get ascending
            'page_obj': page_obj,  # Pass the paginated object to the template
        }

        return render(request, 'archive.html', context)
    except Exception as e:
        # Handle any exceptions, you may want to handle them differently
        return JsonResponse({'error': str(e)}, status=500)  

def report(request):
    try:
        sort_by = request.GET.get('sort', 'created_at')
        descending = request.GET.get('desc')

        # Get all active PDFs for the user
        active_pdfs = PdfTextModel.objects.filter(user=request.user, deleted=False).order_by('-created_at')

        # Get all deleted PDFs for the user
        deleted_pdfs = DeletedPdf.objects.filter(user=request.user).order_by('-deleted_at')

        # Calculate counts
        active_pdf_count = active_pdfs.count()
        deleted_pdf_count = deleted_pdfs.count()

        # Calculate the total count of PDFs (both active and deleted)
        total_obj = active_pdf_count + deleted_pdf_count

        # Get the current date
        current_date = timezone.now()

        context = {
            'active_pdfs': active_pdfs,
            'deleted_pdfs': deleted_pdfs,
            'sort_by': sort_by,
            'ascending': not descending,
            'active_pdf_count': active_pdf_count,
            'deleted_pdf_count': deleted_pdf_count,
            'total_obj': total_obj,
            'current_date': current_date,
        }

        return render(request, 'pdfgo/report.html', context)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    
def generate_pdf_report(request):
    # Create a PDF file
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="report.pdf"'

    # Create a PDF object and add content to it
    p = canvas.Canvas(response)
    p.drawString(100, 750, "My Report")
    # Add more content as needed

    # Close the PDF object cleanly and we're done
    p.showPage()
    p.save()

    return response

def jpgtopdf(request):
    merged_pdf = None

    if request.method == "POST" and request.FILES.getlist('imageFiles'):
        try:
            imageFiles = request.FILES.getlist('imageFiles')

            pdf_merger = PdfMerger()

            for imageFile in imageFiles:
                image = Image.open(imageFile)
                if image.mode == "RGBA":
                    image = image.convert("RGB")

                # Create an in-memory PDF
                pdf_buffer = io.BytesIO()
                image.save(pdf_buffer, "PDF", resolution=100.0)

                # Move the buffer's pointer to the beginning
                pdf_buffer.seek(0)

                # Append the in-memory PDF to the PdfMerger
                pdf_merger.append(pdf_buffer)

            # Create an in-memory merged PDF
            merged_pdf_buffer = io.BytesIO()
            pdf_merger.write(merged_pdf_buffer)

            # Move the buffer's pointer to the beginning
            merged_pdf_buffer.seek(0)

            # Set the merged_pdf variable to the download URL
            merged_pdf = 'data:application/pdf;base64,' + base64.b64encode(merged_pdf_buffer.getvalue()).decode()

            messages.info(request, 'PDF Generated Successfully!')

        except Exception as e:
            print(e)

    return render(request, 'pdfgo/jpg-to-pdf.html', {'merged_pdf': merged_pdf})

# def pdftoword(request):
#     if request.method == "POST" and request.FILES.getlist('pdftowordFiles'):
#         # Get the uploaded PDF file from the request
#         pdf_file = request.FILES['pdftowordFiles']

#         # Create a unique filename for the output file
#         unique_filename = str(uuid.uuid4()) + ".docx"

#         # Save the PDF file to a temporary location
#         temp_pdf_path = 'input.pdf'
#         with open(temp_pdf_path, 'wb') as pdf_temp_file:
#             for chunk in pdf_file.chunks():
#                 pdf_temp_file.write(chunk)

#         # Convert the PDF file to Word (.docx) format
#         temp_docx_path = os.path.join(settings.MEDIA_ROOT, unique_filename)
#         cv = Converter(temp_pdf_path)
#         cv.convert(temp_docx_path)
#         cv.close()

#         # Delete the temporary PDF file
#         os.remove(temp_pdf_path)

#         # Open the generated Word file using python-docx
#         doc = Document(temp_docx_path)

#         # Set the page size to A4 (8.27x11.69 inches)
#         doc.sections[0].page_width = Inches(8.27)
#         doc.sections[0].page_height = Inches(11.69)

#         # Calculate and adjust the margins to center the content
#         margin_left = int((Inches(8.27) - doc.sections[0].left_margin - doc.sections[0].right_margin) / 2)
#         doc.sections[0].left_margin = margin_left
#         doc.sections[0].right_margin = margin_left

#         # Save the changes to the Word document
#         doc.save(temp_docx_path)
#         messages.info(request, 'Word file Generated Successfully!')

#         # Provide the download link for the converted Word file
#         word_file_url = settings.MEDIA_URL + unique_filename

#         return render(request, 'pdfgo/pdf-to-word.html', {'word_file_url': word_file_url})

#     return render(request, 'pdfgo/pdf-to-word.html')

def pdftoword(request):
    if request.method == "POST" and request.FILES.getlist('pdftowordFiles'):
        # Get the uploaded PDF file from the request
        pdf_file = request.FILES['pdftowordFiles']

        # Define a unique temporary PDF file path
        temp_pdf_path = os.path.join(settings.MEDIA_ROOT, 'temp.pdf')

        # Save the PDF file to the temporary path
        with open(temp_pdf_path, 'wb') as pdf_temp_file:
            for chunk in pdf_file.chunks():
                pdf_temp_file.write(chunk)

        # Load the PDF document using PyMuPDF
        pdf_document = fitz.open(temp_pdf_path)

        # Create a new Word document
        doc = Document()

        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)

            # Extract text from the page and add it to the Word document
            text = page.get_text()
            paragraph = doc.add_paragraph(text)

            # Set paragraph style (adjust font and alignment as needed)
            for run in paragraph.runs:
                run.font.size = Pt(12)  # Set the font size to 12 points

            # Process images on the page
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_data = base_image["image"]

                try:
                    # Attempt to open the image data with Pillow
                    image = Image.open(io.BytesIO(image_data))
                    
                    # Convert image_data to RGB if it's in CMYK color mode
                    if base_image["colorspace"] == 4:  # 4 represents CMYK color mode
                        image = image.convert("RGB")

                    # Resize the image if needed (adjust the dimensions as desired)
                    max_width = 6.5 * 72  # Maximum width for an A4 page (6.5 inches)
                    if image.width > max_width:
                        new_width = max_width
                        new_height = (image.height / image.width) * new_width
                        image = image.resize((int(new_width), int(new_height)))

                    # Save the image to a temporary file as PNG
                    img_path = os.path.join(settings.MEDIA_ROOT, f"image_{page_num}_{img_index}.png")
                    image.save(img_path)

                    # Add the image to the Word document with adjusted width
                    doc.add_picture(img_path, width=Inches(image.width / 72))  # Convert width to inches

                    # Clean up the temporary image file
                    os.remove(img_path)

                except Exception as e:
                    # Handle the error (you can log it or skip the image)
                    print(f"Error processing image {img_index}: {str(e)}")

        # Close the PDF document to release the file
        pdf_document.close()

        # Generate a unique filename for the Word document
        unique_filename = str(uuid.uuid4()) + ".docx"

        # Save the Word document in the media directory with the unique filename
        output_docx_path = os.path.join(settings.MEDIA_ROOT, unique_filename)
        doc.save(output_docx_path)

        # Clean up the temporary PDF file
        os.remove(temp_pdf_path)

        # Provide the download link for the converted Word file
        word_file_url = os.path.join(settings.MEDIA_URL, unique_filename)

        return render(request, 'pdfgo/pdf-to-word.html', {'word_file_url': word_file_url})

    return render(request, 'pdfgo/pdf-to-word.html')

def wordtopdf(request):
    if request.method == "POST" and request.FILES.getlist('wordtopdfFiles'):
        # Get the uploaded DOC file from the request
        doc_file = request.FILES['wordtopdfFiles']

        # Load the DOC document using python-docx
        doc = Document(doc_file)

        # Extract the original filename (without extension)
        original_filename = os.path.splitext(doc_file.name)[0]

        # Using reportlab to create a PDF with paragraphs
        pdf_content = []
        styles = getSampleStyleSheet()
        for paragraph in doc.paragraphs:
            # Extract text from the paragraph and add it to the PDF document
            text = paragraph.text
            pdf_content.append(Paragraph(text, styles['Normal']))

        # Create a response object with appropriate content type for PDF
        response = HttpResponse(content_type='application/pdf')
        
        # Set the filename for the response (this will be the default filename when downloaded)
        response['Content-Disposition'] = f'attachment; filename="{original_filename}.pdf"'

        # Using reportlab to create a PDF with paragraphs
        pdf = SimpleDocTemplate(response, pagesize=letter)
        pdf.build(pdf_content)

        return response

    # If the form is not submitted or there's an error, you can redirect or render another page
    return render(request, 'pdfgo/word-to-pdf.html')

def pdftojpg(request):
    if request.method == "POST" and request.FILES.getlist('pdftojpg'):
        # Get the uploaded PDF file from the request
        pdf_file = request.FILES['pdftojpg']

        # Create a unique directory to store the JPEG images
        output_dir = os.path.join(settings.MEDIA_ROOT, str(uuid.uuid4()))
        os.makedirs(output_dir)

        # Save the PDF file to a temporary location
        temp_pdf_path = os.path.join(output_dir, 'input.pdf')
        with open(temp_pdf_path, 'wb') as pdf_temp_file:
            for chunk in pdf_file.chunks():
                pdf_temp_file.write(chunk)

        try:
            # Open the PDF file using PyMuPDF
            pdf_document = fitz.open(temp_pdf_path)
            
            # Specify the output directory for image storage
            output_path = output_dir

            # Convert PDF to JPEG using PyMuPDF
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                image = page.get_pixmap(dpi=200)
                image_path = os.path.join(output_path, f'page_{page_num + 1}.jpg')
                image.save(image_path, "jpeg")

            # Close the PDF document
            pdf_document.close()

            # Create a ZIP archive containing all JPEG images
            zip_filename = f"{uuid.uuid4()}.zip"
            zip_filepath = os.path.join(settings.MEDIA_ROOT, zip_filename)
            with ZipFile(zip_filepath, 'w') as zipf:
                for jpg_file in os.listdir(output_path):
                    jpg_path = os.path.join(output_path, jpg_file)
                    # Skip adding 'input.pdf' to the ZIP archive
                    if jpg_path != temp_pdf_path:
                        zipf.write(jpg_path, os.path.basename(jpg_path))

            # Clean up the temporary PDF and image files
            os.remove(temp_pdf_path)
            for jpg_file in os.listdir(output_path):
                jpg_path = os.path.join(output_path, jpg_file)
                os.remove(jpg_path)
            os.rmdir(output_path)  # Remove the empty directory

            # Provide the download link for the ZIP archive
            zip_file_url = os.path.join(settings.MEDIA_URL, zip_filename)

            return render(request, 'pdfgo/pdf-to-jpg.html', {'zip_file_url': zip_file_url})

        except Exception as e:
            # Handle the exception, for example, show an error message or log the error
            return HttpResponse("Error occurred while converting the PDF to JPEG.")

    return render(request, 'pdfgo/pdf-to-jpg.html')

def pdf_editor(request):
    edited_pdf_url = None

    if request.method == 'POST':
        uploaded_file = request.FILES.get('pdf_file')

        if uploaded_file:
            try:
                # Server-side: Save the uploaded PDF temporarily
                temp_pdf_path = os.path.join(settings.MEDIA_ROOT, 'temp.pdf')
                with open(temp_pdf_path, 'wb') as temp_pdf_file:
                    for chunk in uploaded_file.chunks():
                        temp_pdf_file.write(chunk)

                # Server-side: Provide a URL to the temporarily saved PDF
                temp_pdf_url = os.path.join(settings.MEDIA_URL, 'temp.pdf')

                return render(request, 'pdfgo/pdf_editor.html', {'temp_pdf_url': temp_pdf_url})

            except Exception as e:
                return HttpResponse(f"Error processing PDF: {str(e)}")

    return render(request, 'pdfgo/pdf_editor.html', {'edited_pdf_url': edited_pdf_url})

def save_edited_pdf(request):
    if request.method == 'POST':
        # Client-side: Handle PDF editing using pdf-lib here
        # You can add/edit annotations, text, shapes, etc.

        # After editing, save the edited PDF to a temporary location using PyPDF2
        input_pdf_path = os.path.join(settings.MEDIA_ROOT, 'temp.pdf')
        edited_pdf_path = os.path.join(settings.MEDIA_ROOT, 'edited.pdf')

        try:
            # Open the input PDF
            with open(input_pdf_path, 'rb') as input_pdf:
                pdf_reader = PyPDF2.PdfFileReader(input_pdf)
                pdf_writer = PyPDF2.PdfFileWriter()

                # Perform PDF editing actions here
                # Example: Add a watermark to each page
                watermark = PyPDF2.PdfFileReader(open('path_to_watermark.pdf', 'rb'))
                for page_num in range(pdf_reader.getNumPages()):
                    page = pdf_reader.getPage(page_num)
                    page.mergePage(watermark.getPage(0))
                    pdf_writer.addPage(page)

                # Save the edited PDF to the output file
                with open(edited_pdf_path, 'wb') as edited_pdf:
                    pdf_writer.write(edited_pdf)

            # Provide a URL to the edited PDF
            edited_pdf_url = os.path.join(settings.MEDIA_URL, 'edited.pdf')

            return render(request, 'pdfgo/pdf_editor.html', {'edited_pdf_url': edited_pdf_url})

        except Exception as e:
            return HttpResponse(f"Error editing PDF: {str(e)}")

    return HttpResponse("Invalid request.")


# def passwordprotectpdf(request):
#     pdf_url = None

#     if request.method == "POST" and request.FILES.get('encryptedpdf'):
#         pdf_file = request.FILES['encryptedpdf']
#         out = PdfWriter()

#         # Open the uploaded PDF file with the PdfReader
#         file = PdfReader(pdf_file)

#         # Get number of pages in the original file
#         num = len(file.pages)

#         # Iterate through every page of the original file and add it to our new file.
#         for idx in range(num):
#             # Get the page at index idx
#             page = file.pages[idx]

#             # Add it to the output file
#             out.add_page(page)

#         # Create a variable "password" and store our password in it.
#         password = request.POST.get('pdf_password', '')  # Assumes a form field with name 'pdf_password'

#         # Encrypt the new file with the entered password
#         out.encrypt(password)

#         # Save the encrypted PDF to a new file "myfile_encrypted.pdf"
#         encrypted_file_path = os.path.join('media', 'myfile_encrypted.pdf')
#         with open(encrypted_file_path, "wb") as f:
#             # Write our encrypted PDF to this file
#             out.write(f)

#         # Get the URL for the encrypted PDF
#         pdf_url = '/' + encrypted_file_path.replace('\\', '/')
#         messages.info(request, 'PDF successfully password-protected!')

#     return render(request, 'pdfgo/password-protect-pdf.html', {'pdf_url': pdf_url})

def passwordprotectpdf(request):
    pdf_url = None

    if request.method == "POST" and request.FILES.get('encryptedpdf'):
        pdf_file = request.FILES['encryptedpdf']
        original_pdf_name = pdf_file.name

        # Define the upload directory path
        upload_dir = 'media/'

        # Open the uploaded PDF file with PdfReader
        pdf_reader = PdfReader(pdf_file)

        # Create a PdfWriter to write the encrypted PDF
        pdf_writer = PdfWriter()

        # Copy pages from the original PDF to the PdfWriter
        for page in pdf_reader.pages:
            pdf_writer.add_page(page)

        # Get the password from the form
        password = request.POST.get('pdf_password', '')

        # Encrypt the PDF
        pdf_writer.encrypt(user_pwd=password, owner_pwd=None, use_128bit=True)

        # Construct the filename for the encrypted PDF by adding "encrypted_" as a prefix
        encrypted_pdf_name = f"{original_pdf_name}_encrypted.pdf"
        encrypted_pdf_path = os.path.join(upload_dir, encrypted_pdf_name)

        # Write the encrypted PDF to a new file
        with open(encrypted_pdf_path, "wb") as f:
            pdf_writer.write(f)

        # Get the URL for the encrypted PDF
        pdf_url = '/' + encrypted_pdf_path.replace('\\', '/')

        # Display a success message
        messages.info(request, 'PDF successfully password-protected!')

    return render(request, 'pdfgo/password-protect-pdf.html', {'pdf_url': pdf_url})

def unlockpdf(request):
    decrypted_pdf_url = None

    if request.method == "POST" and request.FILES.get('unlockpdf'):
        out = PdfWriter()

        # Open the uploaded encrypted PDF file with the PdfReader
        encrypted_file = request.FILES['unlockpdf']
        file = PdfReader(encrypted_file)

        # Store the correct password in a variable named "password"
        password = request.POST.get('pdf_password', '')

        # Check if the opened file is actually encrypted
        if file.is_encrypted:

            # If encrypted, decrypt it with the password
            if file.decrypt(password):

                # Now, the file has been unlocked.
                # Iterate through every page of the file
                # and add it to our new file.
                for idx in range(len(file.pages)):
                    # Get the page at index idx
                    page = file.pages[idx]

                    # Add it to the output file
                    out.add_page(page)

                # Save the decrypted PDF to a new file "myfile_decrypted.pdf"
                decrypted_file_path = os.path.join('media', 'myfile_decrypted.pdf')
                with open(decrypted_file_path, "wb") as f:
                    # Write our decrypted PDF to this file
                    out.write(f)

                # Get the URL for the decrypted PDF
                decrypted_pdf_url = '/' + decrypted_file_path.replace('\\', '/')

                # Print success message when done
                messages.info(request, 'File decrypted successfully!')
            else:
                # If the password provided is incorrect
                messages.info(request, 'Incorrect password. File decryption failed.')
        else:
            # If file is not encrypted
            messages.info(request, 'File is not encrypted. No decryption needed.')

    return render(request, 'pdfgo/unlock-pdf.html', {'decrypted_pdf_url': decrypted_pdf_url})

def login_first(request):
    return render(request, 'pdfgo/login_first.html')

@method_decorator(login_required(login_url='login_first'), name='dispatch')
class CreatePdf(CreateView):
    form_class = PdfTextForm
    model = PdfTextModel
    template_name = "pdfgo/create-pdf.html"
    success_message = "Added Successfully"

    def get_success_url(self):
        return reverse('pdf-detail', args=[self.object.pk])

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

@method_decorator(login_required(login_url='login'), name='dispatch')
class PdfEditView(View):
    def get(self, request, pk):
        pdf = get_object_or_404(PdfTextModel, pk=pk)
        form = PdfTextForm(instance=pdf)
        return render(request, 'pdfgo/edit-pdf.html', {'form': form, 'pdf_id': pk})

    def post(self, request, pk):
        pdf = get_object_or_404(PdfTextModel, pk=pk)
        form = PdfTextForm(request.POST, instance=pdf)
        if form.is_valid():
            form.save()
            return redirect('pdf-detail', pk=pdf.pk)
        return render(request, 'pdfgo/edit-pdf.html', {'form': form, 'pdf_id': pk})

@method_decorator(login_required(login_url='login'), name='dispatch')  
class DownloadPdfView(View):
    def get(self, request, pk):
        pdf = get_object_or_404(PdfTextModel, pk=pk)
        template = get_template('pdfgo/pdf-view-download.html')
        context = {'view': pdf}
        html = template.render(context)

        # Create a PDF response
        pdf_response = HttpResponse(content_type='application/pdf')

        # Set the filename based on the name field of the PdfTextModel
        pdf_response['Content-Disposition'] = f'attachment; filename="{pdf.name}.pdf"'

        # Generate PDF using xhtml2pdf
        pisa_status = pisa.CreatePDF(html, dest=pdf_response)

        if pisa_status.err:
            return HttpResponse("Error generating PDF")

        return pdf_response

@method_decorator(login_required(login_url='login'), name='dispatch')
class PdfDetail(DetailView):
    model = PdfTextModel
    context_object_name = 'view'
    template_name = "pdfgo/pdf-view.html"

@method_decorator(login_required(login_url='login'), name='dispatch')
class PdfDeleteView(View):
    def post(self, request, pk):
        pdf = get_object_or_404(PdfTextModel, pk=pk)
        
        # Get the current timestamp when the PDF is deleted
        deleted_at = timezone.now()
        
        # Create a DeletedPdf instance with the original_pdf_id, name, created_at, and deleted_at fields set
        deleted_pdf = DeletedPdf(
            user=request.user,
            original_pdf_id=pdf.id,  # Set the original_pdf_id field to the ID of the PdfTextModel instance
            name=pdf.name,
            created_at=pdf.created_at,  # Include the created timestamp from the original PDF
            deleted_at=deleted_at,  # Include the timestamp when the PDF is deleted
        )
        deleted_pdf.save()
        
        # Delete the PdfTextModel instance
        pdf.delete()

        messages.success(request, f'PDF "{pdf.name}" deleted successfully!')
        return redirect('create-pdf')
